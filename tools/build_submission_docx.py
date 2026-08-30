"""
build_submission_docx.py —— 把最终论文 Markdown 转为 Elsevier 投稿用 Word(.docx)
- 作者信息填好（厉兴/Li Xing + HealthLens），仅邮箱/ORCID 留占位
- 结构化：标题/作者/摘要/章节/表格/参考文献
- 字体 Times New Roman，符合学术投稿惯例
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.join(os.path.dirname(__file__), "..")
SRC = os.path.join(ROOT, "HealthLens_论文_整合医学稳态轴框架_全文.md")
OUT = os.path.join(ROOT, "HealthLens_论文_投稿版.docx")

doc = Document()
# 基础样式
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
for lvl, sz in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12)]:
    st = doc.styles[lvl]
    st.font.name = "Times New Roman"
    st.font.size = Pt(sz)
    st.font.bold = True

def clean(cell: str) -> str:
    return re.sub(r"\*\*", "", cell).strip()

lines = open(SRC, encoding="utf-8").read().split("\n")

i = 0
author_emitted = False
while i < len(lines):
    line = lines[i]
    # 标题
    if line.startswith("# ") and not line.startswith("##"):
        doc.add_paragraph(line[2:].strip(), style="Title")
        i += 1
        continue
    # 作者占位块：渲染为作者/单位段落，不写章节标题
    if line.startswith("## 作者与单位"):
        p = doc.add_paragraph()
        r = p.add_run("Li Xing (厉兴)")
        r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("HealthLens, Independent Research Project, Hangzhou, China",
                          style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Corresponding author: Li Xing ｜ Email: {{CORRESPONDING_EMAIL}} ｜ ORCID: {{ORCID_OPTIONAL}}",
                          style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 跳过该块内的 bullet 行直到下一个 '---' 或 '##'
        i += 1
        while i < len(lines) and not lines[i].startswith("## ") and lines[i].strip() != "---":
            i += 1
        author_emitted = True
        continue
    # 分隔线
    if line.strip() == "---":
        i += 1
        continue
    # 章节标题
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
        i += 1
        continue
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
        i += 1
        continue
    # 引用块
    if line.startswith("> "):
        txt = line[2:].strip()
        # 合并连续引用块
        while i + 1 < len(lines) and lines[i + 1].startswith("> "):
            i += 1
            txt += " " + lines[i][2:].strip()
        para = doc.add_paragraph(txt)
        para.runs[0].italic = True
        para.runs[0].font.size = Pt(10)
        i += 1
        continue
    # 表格
    if line.strip().startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append([clean(c) for c in lines[i].strip().strip("|").split("|")])
            i += 1
        # rows[1] 应为分隔行
        if len(rows) >= 2 and all(set(c) <= set("-: ") for c in rows[1]):
            data = [rows[0]] + rows[2:]
        else:
            data = rows
        if data:
            t = doc.add_table(rows=len(data), cols=len(data[0]))
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(data):
                for ci, val in enumerate(row):
                    cell = t.cell(ri, ci)
                    cell.text = val
                    for pr in cell.paragraphs:
                        for rn in pr.runs:
                            rn.font.size = Pt(9.5)
                            if ri == 0:
                                rn.font.bold = True
        continue
    # 代码块
    if line.strip().startswith("```"):
        i += 1
        code = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code.append(lines[i])
            i += 1
        i += 1  # 跳过结束 ```
        cp = doc.add_paragraph("\n".join(code))
        for rn in cp.runs:
            rn.font.name = "Consolas"
            rn.font.size = Pt(9)
        continue
    # 空行
    if line.strip() == "":
        i += 1
        continue
    # 普通段落
    doc.add_paragraph(line.strip())
    i += 1

doc.save(OUT)
print("Saved:", OUT, "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
