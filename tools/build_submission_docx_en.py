#!/usr/bin/env python3
"""Generic markdown -> submission DOCX converter (English edition)."""
import sys, re, json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = "HealthLens_论文_英文版_全文.md"
OUT = "HealthLens_论文_英文投稿版.docx"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)

AUTHOR = "Li Xing (厉兴)"
AFFIL = "HealthLens, Independent Research Project, Hangzhou, China"

def add_inline(par, text):
    """Handle **bold**, *italic*, `code` inline."""
    # tokenize by backtick first
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*|\\*[^*]+\\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("`") and p.endswith("`"):
            r = par.add_run(p[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif p.startswith("**") and p.endswith("**"):
            r = par.add_run(p[2:-2]); r.bold = True
        elif p.startswith("*") and p.endswith("*"):
            r = par.add_run(p[1:-1]); r.italic = True
        else:
            par.add_run(p)

def flush_table(rows):
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = t.cell(i, j)
            txt = row[j] if j < len(row) else ""
            cell.text = ""
            par = cell.paragraphs[0]
            # header row bold
            if i == 0:
                run = par.add_run(txt.strip())
                run.bold = True
            else:
                add_inline(par, txt.strip())
    doc.add_paragraph()

def parse():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    i = 0
    title_done = False
    while i < len(lines):
        line = lines[i]
        # code fence
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i]); i += 1
            i += 1
            cp = doc.add_paragraph()
            run = cp.add_run("\n".join(code_lines))
            run.font.name = "Consolas"; run.font.size = Pt(9)
            cp.paragraph_format.left_indent = Inches(0.3)
            continue
        # table
        if line.strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # skip separator row like |---|---|
                if set("".join(cells).replace("-", "")) == set() or all(set(c) <= set("-: ") for c in cells):
                    i += 1; continue
                rows.append(cells); i += 1
            if rows:
                flush_table(rows)
            continue
        # blockquote
        if line.lstrip().startswith(">"):
            quotes = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                quotes.append(lines[i].lstrip()[1:].strip()); i += 1
            qp = doc.add_paragraph()
            qp.paragraph_format.left_indent = Inches(0.3)
            run = qp.add_run(" ".join(quotes))
            run.italic = True; run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue
        # title
        if line.startswith("# ") and not title_done:
            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = tp.add_run(line[2:].strip()); run.bold = True; run.font.size = Pt(16)
            title_done = True; i += 1; continue
        # headings
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=1); i += 1; continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=2); i += 1; continue
        # bullet
        if re.match(r"^\s*-\s+", line):
            bp = doc.add_paragraph(style="List Bullet")
            add_inline(bp, re.sub(r"^\s*-\s+", "", line))
            i += 1; continue
        # blank
        if not line.strip():
            i += 1; continue
        # normal paragraph
        p = doc.add_paragraph()
        add_inline(p, line.strip())
        i += 1

parse()
doc.save(OUT)
print("Saved", OUT)
# sanity
d = Document(OUT)
print("paragraphs:", len(d.paragraphs), "tables:", len(d.tables))
